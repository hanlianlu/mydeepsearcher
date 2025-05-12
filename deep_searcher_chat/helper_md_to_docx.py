from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import markdown
from bs4 import BeautifulSoup
import re

try:
    import pypandoc
except ImportError:
    pypandoc = None

def generate_docx(user_input, ai_response):
    """Generate a DOCX file from user input and AI response in Markdown format.
    
    Attempts to use Pandoc for direct Markdown-to-DOCX conversion. Falls back to 
    the manual method (Markdown -> HTML -> DOCX) if Pandoc is unavailable or fails.
    
    Args:
        user_input (str): User input in Markdown format.
        ai_response (str): AI response in Markdown format.
    
    Returns:
        BytesIO: DOCX file content as a BytesIO object.
    """
    if pypandoc is not None:
        try:
            # Combine user input and AI response into a single Markdown string
            full_markdown = "# User Input\n\n" + user_input + "\n\n# AI Response\n\n" + ai_response
            # Convert Markdown to DOCX using Pandoc
            docx_bytes = pypandoc.convert_text(full_markdown, 'docx', format='md')
            doc_io = BytesIO(docx_bytes)
            return doc_io
        except Exception as e:
            print(f"Pandoc conversion failed: {e}")
            # Proceed to fallback method if Pandoc fails

    # Fallback: Manual method using markdown, BeautifulSoup, and python-docx
    doc = Document()
    
    # Add user input section
    doc.add_heading("User Input", level=1)
    user_html = markdown.markdown(user_input, extensions=['tables'])
    user_soup = BeautifulSoup(user_html, 'html.parser')
    add_html_to_docx(user_soup, doc)
    
    # Add AI response section
    doc.add_heading("AI Response", level=1)
    ai_html = markdown.markdown(ai_response, extensions=['tables'])
    ai_soup = BeautifulSoup(ai_html, 'html.parser')
    add_html_to_docx(ai_soup, doc)
    
    # Save document to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def add_html_to_docx(soup, doc, list_level=0, is_ordered=False):
    """Add parsed HTML content to a DOCX document.
    
    Args:
        soup: BeautifulSoup object containing parsed HTML.
        doc: python-docx Document object to add content to.
        list_level (int): Current list nesting level.
        is_ordered (bool): Whether the list is ordered (True) or unordered (False).
    """
    for element in soup.children:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            p = doc.add_heading(element.text.strip(), level=level)
            p.paragraph_format.left_indent = Pt(0)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif element.name == 'p':
            text = element.text.strip()
            indent_level = detect_indent_level(text)
            if indent_level > 0:
                p = doc.add_paragraph(text.lstrip(), style='ListBullet')
                set_list_level(p, indent_level - 1, is_ordered=False)
            else:
                p = doc.add_paragraph(text)
                p.paragraph_format.left_indent = Pt(0)
            add_inline_elements(element, p)
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.text.strip(), style='ListBullet')
                set_list_level(p, list_level, is_ordered=False)
                add_inline_elements(li, p)
                # Handle nested lists
                nested_ul = li.find('ul', recursive=False)
                if nested_ul:
                    add_html_to_docx(nested_ul, doc, list_level + 1, is_ordered=False)
                nested_ol = li.find('ol', recursive=False)
                if nested_ol:
                    add_html_to_docx(nested_ol, doc, list_level + 1, is_ordered=True)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.text.strip(), style='ListNumber')
                set_list_level(p, list_level, is_ordered=True)
                add_inline_elements(li, p)
                # Handle nested lists
                nested_ul = li.find('ul', recursive=False)
                if nested_ul:
                    add_html_to_docx(nested_ul, doc, list_level + 1, is_ordered=False)
                nested_ol = li.find('ol', recursive=False)
                if nested_ol:
                    add_html_to_docx(nested_ol, doc, list_level + 1, is_ordered=True)
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['th', 'td'])
                    for j, cell in enumerate(cells):
                        doc_cell = table.rows[i].cells[j]
                        p = doc_cell.add_paragraph()
                        add_cell_content(cell, p)
                        align = cell.get('align')
                        if align:
                            if align == 'left':
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            elif align == 'center':
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            elif align == 'right':
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        if cell.name == 'th':
                            for run in p.runs:
                                run.bold = True
        elif element.name == 'pre':
            code = element.find('code')
            if code:
                p = doc.add_paragraph()
                run = p.add_run(code.text)
                run.font.name = 'Courier New'
        else:
            pass  # Skip unhandled elements

def detect_indent_level(text):
    """Detect the indentation level of a text string based on leading spaces.
    
    Args:
        text (str): Text to analyze.
    
    Returns:
        int: Number of indentation levels (2 spaces = 1 level).
    """
    match = re.match(r'(\s+)', text)
    if match:
        spaces = len(match.group(1))
        return spaces // 2  # 2 spaces per level for compatibility
    return 0

def set_list_level(paragraph, level, is_ordered):
    """Set the list level and style for a paragraph.
    
    Args:
        paragraph: python-docx Paragraph object.
        level (int): List nesting level.
        is_ordered (bool): True for numbered lists, False for bulleted lists.
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1' if is_ordered else '2')
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)
    paragraph.paragraph_format.left_indent = Pt(18 * level)  # Consistent indent
    paragraph.paragraph_format.first_line_indent = Pt(-6)    # Hanging indent

def add_cell_content(cell, paragraph):
    """Add content from an HTML table cell to a DOCX paragraph.
    
    Args:
        cell: BeautifulSoup cell element (th or td).
        paragraph: python-docx Paragraph object to add content to.
    """
    for child in cell.children:
        if child.name == 'strong':
            paragraph.add_run(child.text).bold = True
        elif child.name == 'em':
            paragraph.add_run(child.text).italic = True
        elif child.name == 'code':
            paragraph.add_run(child.text).font.name = 'Courier New'
        elif isinstance(child, str):
            paragraph.add_run(child)
        else:
            paragraph.add_run(child.text)

def add_inline_elements(element, paragraph):
    """Add inline HTML elements (e.g., bold, italic, code) to a DOCX paragraph.
    
    Args:
        element: BeautifulSoup element containing inline content.
        paragraph: python-docx Paragraph object to add content to.
    """
    for child in element.children:
        if child.name == 'strong':
            paragraph.add_run(child.text).bold = True
        elif child.name == 'em':
            paragraph.add_run(child.text).italic = True
        elif child.name == 'code':
            paragraph.add_run(child.text).font.name = 'Courier New'
        elif isinstance(child, str):
            paragraph.add_run(child.strip())
        else:
            paragraph.add_run(child.text.strip())