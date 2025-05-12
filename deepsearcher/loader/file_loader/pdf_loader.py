from __future__ import annotations

import io
import logging
import os
import time
from concurrent.futures import ThreadPoolExecutor
from typing import List, Optional

import fitz  # PyMuPDF
from PIL import Image
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_community.document_loaders import PDFPlumberLoader
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from azure.core.exceptions import HttpResponseError
from docx import Document as DocxDocument

# Project loaders
from deepsearcher.loader.file_loader.base import BaseLoader
from deepsearcher.loader.file_loader.excel_loader import ExcelRAGLoader
from deepsearcher.loader.file_loader.pptx_loader import AdvancedPPTXLoader
from deepsearcher.loader.file_loader.docling_loader import DoclingLoader

logger = logging.getLogger(__name__)

# Best Practice: Test your documents in Document Intelligence Studio to verify the output before running this script.
# See https://learn.microsoft.com/en-us/azure/ai-services/document-intelligence/studio-overview

class PDFLoader(BaseLoader):
    # ------------------------------------------------------------------ #
    def __init__(self, **kwargs):
        logging.basicConfig(level=logging.INFO)
        self.logger = logger

        self.excel_loader = ExcelRAGLoader()
        self.pptx_loader = AdvancedPPTXLoader(ocr_images=True)
        self.docling_loader = DoclingLoader()

        # Azure setup
        load_dotenv()
        ep = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
        key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")
        self.azure_client = (
            DocumentIntelligenceClient(endpoint=ep, credential=AzureKeyCredential(key))
            if ep and key else None
        )
        if not self.azure_client:
            self.logger.warning("Azure credentials not found – Azure path disabled")

        # EasyOCR reader (may be None)
        self.ocr_reader = self._init_easyocr()

    # ------------------------------------------------------------------ #
    @staticmethod
    def _init_easyocr() -> Optional[object]:
        """
        Import EasyOCR lazily and return a Reader instance, or None on failure.
        """
        try:
            import torch
            from easyocr import Reader
        except ImportError as e:
            logger.warning("easyocr/torch import failed (%s) – OCR disabled", e)
            return None

        use_gpu = torch.cuda.is_available()
        for attempt_gpu in (use_gpu, False) if use_gpu else (False,):
            try:
                reader = Reader(lang_list=["en"], gpu=attempt_gpu, verbose=False)
                return reader
            except Exception as e:
                logger.warning("EasyOCR init failed on %s (%s)", "GPU" if attempt_gpu else "CPU", e)
        logger.warning("EasyOCR unavailable – OCR disabled")
        return None

    # ------------------------------------------------------------------ #
    @staticmethod
    def _page_to_image(path: str, page_idx: int) -> Image.Image:
        with fitz.open(path) as pdf:
            pix = pdf.load_page(page_idx).get_pixmap(dpi=300)
        return Image.open(io.BytesIO(pix.tobytes())).convert("RGB")

    def _ocr_image(self, img: Image.Image) -> str:
        if not self.ocr_reader:
            return ""
        try:
            lines = self.ocr_reader.readtext(img, detail=0, paragraph=True)
            return "\n".join(lines)
        except Exception as e:  # noqa: BLE001
            self.logger.warning("EasyOCR runtime error: %s", e)
            return ""

    # ------------------------------------------------------------------ #
    def _load_pdfplumber(self, path: str) -> List[Document]:
        docs = PDFPlumberLoader(path).load()

        def fix(doc: Document) -> Document:
            if doc.page_content.strip():
                return doc
            idx = doc.metadata.get("page", 1) - 1
            img = self._page_to_image(path, idx)
            doc.page_content = self._ocr_image(img) or "No extractable text."
            return doc

        with ThreadPoolExecutor() as pool:
            return list(pool.map(fix, docs))

    # ---------- Azure -------------------------------------------------- #
    def _process_with_azure(self, path: str) -> List[Document]:
        if not self.azure_client:
            raise RuntimeError("Azure client not initialised")

        result = None
        for attempt in range(3):
            try:
                with open(path, "rb") as fh:
                    poller = self.azure_client.begin_analyze_document("prebuilt-layout", fh)
                    result = poller.result()
                self.logger.info(f"Processed {len(result.pages)} pages with Azure for {os.path.basename(path)}.")
                break
            except HttpResponseError as e:
                if attempt < 2:
                    self.logger.warning(f"Retrying due to HTTP error for {os.path.basename(path)}: {e}")
                    time.sleep(2 ** attempt)
                else:
                    self.logger.warning(f"Azure analysis failed after retries for {os.path.basename(path)}: {e}")
                    result = None
            except Exception as e:
                self.logger.warning(f"Azure analysis failed for {os.path.basename(path)}: {type(e).__name__}: {str(e)}")
                result = None

        if result is None:
            return self._load_pdfplumber(path)

        docs: List[Document] = []
        for p in getattr(result, "paragraphs", []) or []:
            docs.append(
                Document(
                    page_content=p.content,
                    metadata={
                        "source": path,
                        "page": p.bounding_regions[0].page_number if p.bounding_regions else None,
                        "type": "paragraph",
                    },
                )
            )
        for tab in getattr(result, "tables", []) or []:
            if not hasattr(tab, "rows") or not tab.rows:
                self.logger.warning("Table without rows or missing rows attribute detected for %s.", os.path.basename(path))
                continue
            headers = [c.content for c in tab.rows[0].cells] if tab.rows else []
            rows = tab.rows[1:] if len(tab.rows) > 1 else []
            if not headers and rows:
                headers = [f"Column {i+1}" for i in range(len(rows[0].cells))]
            if not rows:
                self.logger.warning("Table has no data rows, only headers or empty for %s.", os.path.basename(path))
                continue
            lines = [
                ", ".join(f"{headers[j]}: {cell.content}" for j, cell in enumerate(r.cells))
                for r in rows
            ]
            docs.append(
                Document(
                    page_content="\n".join(lines),
                    metadata={
                        "source": path,
                        "page": tab.bounding_regions[0].page_number if tab.bounding_regions else None,
                        "type": "table",
                    },
                )
            )
        if not docs:
            self.logger.info("No paragraphs or tables extracted by Azure for %s, falling back to PDFPlumber.", os.path.basename(path))
            return self._load_pdfplumber(path)
        return docs

    # ---------- Dispatcher -------------------------------------------- #
    def load_file(self, path: str) -> List[Document]:
        ext = os.path.splitext(path)[1].lower().lstrip(".")
        if ext == "pdf":
            return self._load_pdf(path)
        return self._load_other(path, ext)

    def _load_pdf(self, path: str) -> List[Document]:
        # Note: For best results with Azure Document Intelligence, ensure the PDF is of high quality, with clear text and minimal noise.
        if self.azure_client:
            self.logger.info("Trying Azure for %s", os.path.basename(path))
            try:
                azure_docs = self._process_with_azure(path)
                if azure_docs:
                    return azure_docs
                else:
                    self.logger.warning("Azure returned no documents for %s, falling back to Docling.", os.path.basename(path))
            except Exception as e:
                self.logger.warning("Azure path failed for %s: %s", os.path.basename(path), e)

        self.logger.info("Trying Docling for %s", os.path.basename(path))
        try:
            d = self.docling_loader.load_file(path)
            if d:
                return d
        except Exception as e:
            self.logger.warning("Docling failed for %s: %s", os.path.basename(path), e)

        self.logger.info("Trying PDFPlumber+OCR for %s", os.path.basename(path))
        return self._load_pdfplumber(path)

    # ---------- Non‑PDF helpers --------------------------------------- #
    def _load_other(self, path: str, ext: str) -> List[Document]:
        if ext in self.docling_loader.supported_file_types:
            try:
                return self.docling_loader.load_file(path)
            except Exception as e:
                self.logger.warning("Docling failed for %s: %s", os.path.basename(path), e)

        if ext == "docx":
            return self._load_docx(path)
        if ext in {"txt", "md"}:
            return self._load_text(path)
        if ext in {"xlsx", "xls"}:
            return self.excel_loader.load_file(path)
        if ext == "pptx":
            return self.pptx_loader.load_pptx(path)

        logger.warning("Unsupported extension: %s", ext)
        return []

    # ---------- Simple loaders ---------------------------------------- #
    @staticmethod
    def _load_text(path: str) -> List[Document]:
        try:
            with open(path, "r", encoding="utf-8") as fh:
                return [Document(page_content=fh.read(), metadata={"source": path})]
        except Exception as e:  # noqa: BLE001
            logger.error("TXT/MD read error %s: %s", path, e)
            return []

    @staticmethod
    def _load_docx(path: str) -> List[Document]:
        try:
            docx = DocxDocument(path)
        except Exception as e:  # noqa: BLE001
            logger.error("DOCX read error %s: %s", path, e)
            return []
        docs: List[Document] = []
        docs.extend(
            Document(page_content=p.text, metadata={"source": path, "type": "paragraph"})
            for p in docx.paragraphs if p.text.strip()
        )
        docs.extend(
            Document(
                page_content="\n".join(",".join(c.text for c in r.cells) for r in t.rows),
                metadata={"source": path, "type": "table"},
            )
            for t in docx.tables
        )
        return docs

    # ---------- Directory loader -------------------------------------- #
    def load_directory(self, dir_path: str) -> List[Document]:
        collected: List[Document] = []
        for root, _, files in os.walk(dir_path):
            for f in files:
                if f.split(".")[-1].lower() in self.supported_file_types:
                    collected.extend(self.load_file(os.path.join(root, f)))
        logger.info("Loaded %d docs from %s", len(collected), dir_path)
        return collected

    # ---------- Supported types --------------------------------------- #
    @property
    def supported_file_types(self) -> List[str]:
        base = ["txt", "md", "xls", "xlsx", "pptx", "pdf", "docx"]
        return list(set(base + self.docling_loader.supported_file_types))